#!/usr/bin/env python3
"""
Converts resume.md to a nicely formatted .docx matching the HTML style
Requires: pip install python-docx pyyaml
"""

import re
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# Colors matching the HTML
NAVY = RGBColor(26, 54, 93)
GOLD = RGBColor(184, 134, 11)
DARK_GRAY = RGBColor(45, 55, 72)
MEDIUM_GRAY = RGBColor(74, 85, 104)


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_bottom_border(paragraph, color_hex="B8860B", width="12"):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), width)
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def parse_markdown(md_path: Path) -> tuple[dict, str]:
    """Parse YAML front matter and markdown content."""
    text = md_path.read_text()
    match = re.match(r'^---\n(.*?)\n---\n(.*)$', text, re.DOTALL)
    if match:
        front_matter = yaml.safe_load(match.group(1))
        content = match.group(2)
        return front_matter, content
    return {}, text


def create_resume_docx(md_path: Path, output_path: Path):
    """Create a formatted DOCX from the markdown resume."""
    front_matter, content = parse_markdown(md_path)

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # === HEADER SECTION (Navy background) ===
    # Use a table to create the navy background header
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    set_cell_shading(header_cell, "1A365D")  # Navy

    # Name
    name_para = header_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_para, before=200, after=100)
    name_run = name_para.add_run(front_matter.get('name', '').upper())
    name_run.bold = True
    name_run.font.size = Pt(28)
    name_run.font.color.rgb = RGBColor(255, 255, 255)  # White
    name_run.font.name = 'Arial'

    # Title
    if front_matter.get('title'):
        title_para = header_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(title_para, before=0, after=200)
        title_run = title_para.add_run(front_matter['title'])
        title_run.italic = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = GOLD
        title_run.font.name = 'Georgia'

    # === CONTACT BAR ===
    contact_parts = []
    if front_matter.get('location'):
        contact_parts.append(front_matter['location'])
    if front_matter.get('email'):
        contact_parts.append(front_matter['email'])
    if front_matter.get('linkedin'):
        contact_parts.append("LinkedIn")
    if front_matter.get('website'):
        contact_parts.append("Website")
    if front_matter.get('scholar'):
        contact_parts.append("Google Scholar")

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(contact_para, before=150, after=150)
        add_bottom_border(contact_para, "B8860B", "18")  # Gold border
        contact_run = contact_para.add_run('  \u25c6  '.join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.name = 'Arial'
        contact_run.font.color.rgb = MEDIUM_GRAY

    # === MAIN CONTENT ===
    lines = content.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # H2 - Section headers (EXECUTIVE SUMMARY, EXPERIENCE, etc.)
        if line.startswith('## '):
            header_text = line[3:].strip()
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=300, after=100)
            run = para.add_run(header_text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            run.font.name = 'Arial'
            # Add gold underline
            underline_para = doc.add_paragraph()
            set_paragraph_spacing(underline_para, before=0, after=150)
            add_bottom_border(underline_para, "B8860B", "12")

        # H3 - Company/School names
        elif line.startswith('### '):
            header_text = line[4:].strip()
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=250, after=50)
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = NAVY
            run.font.name = 'Arial'

        # H4 - Role subsections (Data & Infrastructure, etc.)
        elif line.startswith('#### '):
            header_text = line[5:].strip()
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=180, after=80)
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = GOLD
            run.font.name = 'Arial'

        # List items
        elif line.startswith('- '):
            item_text = line[2:].strip()
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=0, after=60)
            # Add bullet character
            bullet_run = para.add_run('\u25b8  ')  # Triangle bullet like HTML
            bullet_run.font.color.rgb = GOLD
            bullet_run.font.size = Pt(10)
            add_formatted_text(para, item_text)

        # Bold lines (role titles like "Research Scientist | Technical Lead")
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=0, after=0)
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Georgia'
            run.font.color.rgb = DARK_GRAY

        # Italic lines (dates)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            text = line[1:-1]
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=0, after=100)
            run = para.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = MEDIUM_GRAY
            run.font.name = 'Georgia'

        # Regular paragraphs
        elif line.strip():
            para = doc.add_paragraph()
            set_paragraph_spacing(para, before=0, after=120)
            add_formatted_text(para, line)

        i += 1

    doc.save(output_path)
    print(f"Successfully created: {output_path}")


def add_formatted_text(para, text):
    """Add text with bold/italic markdown formatting."""
    # Pattern for **bold**, *italic*, and [links](url)
    pattern = r'(\*\*.*?\*\*|\*[^*]+?\*|\[.*?\]\(.*?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = NAVY
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and '](' in part:
            # Link: [text](url) - just show text
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                run = para.add_run(link_match.group(1))
                run.font.color.rgb = NAVY
                run.underline = True
        else:
            run = para.add_run(part)
            run.font.name = 'Georgia'
            run.font.size = Pt(11)


if __name__ == "__main__":
    script_dir = Path(__file__).parent
    md_path = script_dir / "resume.md"
    output_path = script_dir / "resume.docx"
    create_resume_docx(md_path, output_path)
